from docx import Document
from docx.shared import Inches
import pyttsx3
import speech_recognition as sr

def speak(text):
    pyttsx3.speak(text)

#recognizer init
recognizer_instance = sr.Recognizer() 

def transcribe():
    with sr.Microphone() as source:
        recognizer_instance.adjust_for_ambient_noise(source)
        print("...")
        audio = recognizer_instance.listen(source)
    try:
        text = recognizer_instance.recognize_google(audio, language="it-IT")
        return text
    except Exception as e:
        print(e)

document = Document()

#profile picture
#document.add_picture(
#    'profile_picture.jpg',
#    width=Inches(2.0)
#)


#name, phone number and email
speak('Ciao, qual è il tuo nome?')
firstname = transcribe()
speak(firstname +', qual è il tuo cognome? ')
lastname = transcribe()

name=firstname + ' ' + lastname

speak('Qual''è il tuo numero di telefono?')
phone_number = input('Qual''è il tuo numero di telefono? ')

speak('Qual''è la tua mail?')
email = transcribe()

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

#about me
document.add_heading('Chi sono')
speak('Parlami di te')
about_me = transcribe()
document.add_paragraph(about_me)

#work experience
document.add_heading('Esperienze lavorative')
p = document.add_paragraph()

speak('In quale azienda hai lavorato?')
company = transcribe()
speak('Da quando hai lavorato in '+company+'?')
from_date = input('Dalla data ')
speak('Fino a quando?')
to_date = input('Alla data ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

speak('Descrivi la tua esperienza in ' + company)
experience_details = transcribe()
p.add_run(experience_details)

#skills
document.add_heading('Competenze')
speak('Quali competenze hai?')
skill = transcribe()
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    speak('Vuoi inserire un''altra competenza? Rispondere SI o NO')
    q = transcribe()
    if q.lower() == 'si' :
        speak('Quali competenze hai?')
        skill = transcribe()
        p = document.add_paragraph(skill)
        p.style= 'List Bullet'
    else:
        break

#more experiences
while True:
    speak('Hai altre esperienze lavorative? Rispondi SI o NO')
    has_more_exp = transcribe()
    if has_more_exp.lower == 'si':
        speak('In quale azienda hai lavorato?')
        company = transcribe()
        speak('Da quando hai lavorato in '+company+'?')
        from_date = input('Dalla data ')
        speak('Fino a quando?')
        to_date = input('Alla data ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        speak('Descrivi la tua esperienza in '+company)
        experience_details = transcribe()
        p.add_run(experience_details)
    else:
        break

#footer
#section = document.section[0]
#footer = section.footer
#p = footer.paragraph[0]
#p.text = "Cv generated with python"
#
#doc saving
document.save('cv.docx')
